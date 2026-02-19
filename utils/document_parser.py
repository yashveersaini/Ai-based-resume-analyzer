import fitz  # PyMuPDF for PDF
from docx import Document  # python-docx for DOCX
import io
import os


SUPPORTED_EXTENSIONS = {'.pdf', '.docx'}


def get_file_extension(filename: str) -> str:
    """Extract file extension in lowercase with dot (e.g., '.pdf')."""
    return os.path.splitext(filename.lower())[1]


def extract_text_from_pdf(file_stream) -> str:
    """Extract text from PDF using PyMuPDF."""
    doc = fitz.open(stream=file_stream.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text.strip()



from docx import Document

def extract_text_from_docx(file_stream) -> str:
    doc = Document(file_stream)
    full_text = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    full_text.append(text)

    # Remove duplicates while preserving order
    seen = set()
    unique_text = []
    for line in full_text:
        if line not in seen:
            seen.add(line)
            unique_text.append(line)

    return "\n".join(unique_text)


def extract_text_from_document(file_obj, filename: str) -> str:
    """
    Main function: extract text from PDF or DOCX files.
    
    Args:
        file_obj: File object from Flask request.files (has .read(), .seek())
        filename: Original filename for extension detection
    
    Returns:
        Extracted text string
    
    Raises:
        ValueError: If file format is unsupported
        Exception: If extraction fails
    """
    ext = get_file_extension(filename)
    
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file format: {ext}. "
            f"Only PDF and DOCX files are accepted."
        )
    
    # Reset stream position to start
    file_obj.seek(0)
    
    try:
        if ext == '.pdf':
            return extract_text_from_pdf(file_obj)
        elif ext == '.docx':
            return extract_text_from_docx(file_obj)
        else:
            # Should never reach here due to validation above
            raise ValueError(f"Unhandled extension: {ext}")
    
    except Exception as e:
        raise Exception(f"Failed to extract text from {ext} file: {str(e)}")


def is_supported_format(filename: str) -> bool:
    """Check if the file format is supported."""
    ext = get_file_extension(filename)
    return ext in SUPPORTED_EXTENSIONS


def get_supported_formats_string() -> str:
    """Return a user-friendly string of supported formats."""
    return "PDF, DOCX"


if __name__ == "__main__":
    path = "D:\\yashveer_resume.docx"
    with open(path, "rb") as f:
        result = extract_text_from_docx(f)

    print(result)
